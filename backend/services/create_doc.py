from docx import Document
from fpdf import FPDF
import io

def create_word_doc(disease:str, article:str):
    doc = Document()
    doc.add_heading(f"{disease} - Disease Overview", 0)

    for line in article.split("\n"):
        doc.add_paragraph(line)
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

def create_pdf_doc(disease:str, article:str):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.multi_cell(0, 7, f"{disease} - Disease Overview\n\n{article}")

    pdf_stream = io.BytesIO()
    pdf.output(pdf_stream)
    pdf_stream.seek(0)
    return pdf_stream
